from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"D:\work\15_지선차단기")
OUT = BASE / "09_보고서" / "SEMS_타기업_가격특징_비교보고서_v0.1.docx"
USD_KRW = 1500
EUR_KRW = 1737.79
THB_KRW = 46.89758467143435


def won(value):
    return f"{int(round(value)):,}원"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.style = doc.styles["Heading 1"]
    elif level == 2:
        p.style = doc.styles["Heading 2"]
    else:
        p.style = doc.styles["Heading 3"]
    r = p.add_run(text)
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        r.font.size = Pt(10.5)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].width = Cm(widths_cm[i])
        set_cell_text(hdr[i], head, bold=True, size=10.5, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E78")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Cm(widths_cm[i])
            set_cell_text(cells[i], value, size=9.5)
    return table


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.8)
sec.bottom_margin = Cm(1.6)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.6)

styles = doc.styles
for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
    style = styles[style_name]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("SEMS 타기업 가격특징 비교보고서")
r.bold = True
r.font.name = "맑은 고딕"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(31, 78, 121)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("작성일: 2026-04-14 | 기준: 공식 자료 + 공개 유통가 + 구독형 가격모델 확인")
r.font.name = "맑은 고딕"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
r.font.size = Pt(10)

doc.add_paragraph("")

add_heading(doc, "1. 보고서 목적", 1)
doc.add_paragraph(
    "본 보고서는 Smart Load(기존 지선차단기)와 직접 또는 간접 경쟁하는 주요 기업들의 가격 특성과 판매 구조를 비교해, "
    "우리 제품의 가격 포지셔닝과 제안 방식을 정리하기 위한 문서이다."
)

add_heading(doc, "2. 핵심 요약", 1)
add_bullets(doc, [
    "글로벌 경쟁사는 공개 단가 판매보다 프로젝트 견적형, 유통 채널형, SaaS 구독형으로 가격 구조가 나뉜다.",
    "분기회로 모니터링 장비는 공개 유통가 기준 수천 달러급 제품이 많아 초기 도입비가 높다.",
    "상위 플랫폼은 하드웨어와 별도로 연간 구독료 또는 사이트별 과금 구조를 가진다.",
    "우리 Smart Load는 말단 실행형 장치로서 장비 단가와 구축비를 분리해 제안하면 가격 경쟁력을 만들 수 있다.",
])

add_heading(doc, "3. 조사 기준 및 해석 주의사항", 1)
add_bullets(doc, [
    "가격은 2026-04-14 기준 공개적으로 확인 가능한 정보만 사용했다.",
    f"환율 표기 기준은 USD 1달러 = {USD_KRW:,}원으로 고정 적용했다.",
    f"EUR 환산은 1유로 = {EUR_KRW:,.2f}원, THB 환산은 1바트 = {THB_KRW:,.2f}원 기준으로 병기했다.",
    "유통가와 공식 견적가는 다를 수 있으므로, 본 문서는 절대가격보다 가격 구조와 포지셔닝을 보는 용도로 사용해야 한다.",
    "일부 상위 플랫폼은 가격을 공개하지 않고 연간 구독 또는 프로젝트별 견적만 제공한다.",
])

add_heading(doc, "4. 가격 특징 종합 비교표", 1)
headers = [
    "회사",
    "대표 제품",
    "공개 가격/모델",
    "가격 구조 특징",
    "우리에게 주는 시사점",
]
rows = [
    [
        "Schneider Electric",
        "PowerTag Energy\n(A9MEM1520 예시)",
        "Newark 공개가\n1개 $224.48 (336,720원)\n5개 $212.70 (319,050원)\n10개 $204.20 (306,300원)\n25개 $195.69 (293,535원)\n50개 $180.10 (270,150원)",
        "센서 단품 공개가 존재.\n회로당 센서형 구조라 수량이 늘면 총비용이 빠르게 증가.\n상위 플랫폼 비용은 별도.",
        "회로 수가 많아질수록 센서 누적비용이 커진다.\n우리는 말단 장치 단가 + 시스템 구축비를 분리해 제안할 필요가 있다.",
    ],
    [
        "Siemens",
        "Sentron Powermind",
        f"공식 Product Sheet 기준\n120 EUR / year\n약 {won(120 * EUR_KRW)} / year\n구독형",
        "SaaS 애플리케이션 구독형.\n12개월 또는 36개월 term.\n최종 가격은 주문서 기준.",
        "상위 분석 기능은 구독형이 일반적이다.\n우리도 향후 서버/분석 기능은 유지비 모델을 분리할 수 있다.",
    ],
    [
        "Eaton",
        "PXBCM-MB-BASIC",
        f"PLATT 공개가\n$2,110.92 / EA\n약 {won(2110.92 * USD_KRW)} / EA",
        "분기회로 모니터링용 베이스 장비가 수천 달러급.\n모듈, CT, 설치비를 더하면 총도입비가 더 상승 가능.",
        "모니터링만 해도 초기 도입비가 높다.\n우리는 차단 기능까지 포함한 말단 장치로 ROI 메시지를 강화해야 한다.",
    ],
    [
        "ABB",
        "Emax 2",
        f"공식 태국 2026 가격표 검색 스니펫 기준\n3P Fixed Type None: 115,010 THB\n약 {won(115010 * THB_KRW)}",
        "배전반급 스마트 차단기라 장치 자체 가격대가 높다.\n프로젝트/구성별 편차가 큼.\n고급 옵션 추가 시 비용 상승.",
        "ABB는 상위 배전단 스마트 차단기 포지션이다.\n우리는 말단 회로용 장치로 가격 접근성을 강조할 수 있다.",
    ],
    [
        "Leviton",
        "VerifEye 71D48",
        f"Lighting Supply 공개가\n$4,452.90\n약 {won(4452.90 * USD_KRW)}",
        "48회로 모니터링 일체형 장비.\n다회로 모니터링이라 초기 장비가는 높지만 회로당 단가로 분산 가능.",
        "우리는 다회로 집합형이 아니라 회로 단위 실행형 장치다.\n설치 방식과 회로 확장 구조를 어떻게 설명할지 중요하다.",
    ],
    [
        "Honeywell",
        "Forge for Buildings",
        "공개 정액은 제한적.\n공식 페이지 기준 annual subscription,\nBMS point 수 기반 tier,\n필요 하드웨어 별도",
        "전형적인 SaaS + 하드웨어 별도 + 서비스형 구조.\n사이트별 과금이라 대형 포트폴리오에 유리.",
        "우리는 상위 SaaS를 지금 당장 따라가기보다 현장 데이터 원천 장치로 먼저 자리 잡는 것이 현실적이다.",
    ],
    [
        "LS ELECTRIC",
        "K-EMS / Smart Energy",
        "공개 단가 미확인\n대부분 프로젝트 견적형",
        "국내 대기업형 솔루션 구조.\n시스템 단위 제안이 많고, 제품 단일 가격보다 프로젝트 견적 비중이 큼.",
        "국내 고객 대상으론 빠른 PoC와 현장 맞춤형 단가 구조가 차별화 포인트가 된다.",
    ],
    [
        "HD현대일렉트릭",
        "INTEGRICT / BEMS",
        "공개 단가 미확인\n프로젝트 견적형",
        "전력기기 + 솔루션 패키지형.\n설계/구축/운영을 함께 보는 구조.",
        "우리는 소규모 현장에서도 시작 가능한 단계형 제안이 필요하다.",
    ],
]
add_table(doc, headers, rows, [2.4, 3.0, 4.2, 4.6, 4.8])

doc.add_paragraph("")
add_heading(doc, "5. 가격 상세 해설", 1)

vendor_details = [
    (
        "5.1 Schneider Electric",
        [
            f"PowerTag Energy는 공개 유통가가 확인되는 드문 케이스다. Newark 기준 A9MEM1520은 1개 $224.48({won(224.48 * USD_KRW)}), 50개 구간은 $180.10({won(180.10 * USD_KRW)})까지 내려간다.",
            "즉 회로 수가 많아질수록 총센서비가 커지지만, 동시에 수량 할인 구조가 존재한다.",
            "가격 구조상 회로 단위 센서형이므로 소규모 설치는 유연하지만 대규모로 갈수록 누적 장비비가 커질 수 있다.",
        ],
    ),
    (
        "5.2 Siemens",
        [
            f"Sentron Powermind는 하드웨어 일시구매가 아니라 공식 Product Sheet 기준 연 120유로(약 {won(120 * EUR_KRW)})의 애플리케이션 구독형이다.",
            "이는 상위 분석/시각화 기능이 반복매출 구조로 설계돼 있음을 의미한다.",
            "우리도 향후 서버 기능을 붙일 경우 장치 판매와 데이터 서비스 과금을 분리하는 모델을 고려할 수 있다.",
        ],
    ),
    (
        "5.3 Eaton",
        [
            f"PLATT 공개가 기준 PXBCM-MB-BASIC은 $2,110.92/EA, 한화 약 {won(2110.92 * USD_KRW)} 수준이다.",
            "여기서 중요한 점은 이 가격이 전체 프로젝트 총비용이 아니라 베이스 장비 가격이라는 점이다.",
            "실제 도입 시 CT, 설치 공사, 통신 연동 비용이 추가될 수 있으므로 총소유비용(TCO)은 더 커질 가능성이 높다.",
        ],
    ),
    (
        "5.4 ABB",
        [
            "ABB Emax 2는 배전반급 스마트 차단기 포지션이라 구조적으로 고가 영역이다.",
            f"공식 태국 2026 가격표 검색 결과에서는 3P Fixed Type None 항목이 115,010 THB로 노출되며, 한화로는 약 {won(115010 * THB_KRW)} 수준이다.",
            "이는 말단 회로 단위 장치가 아니라 상위 배전 차단기 시장이라는 점을 보여준다.",
        ],
    ),
    (
        "5.5 Leviton",
        [
            f"Lighting Supply 공개가 기준 71D48은 $4,452.90이며, 한화 약 {won(4452.90 * USD_KRW)} 수준이다.",
            "48회로 모니터링 일체형이라 초기 장비가는 높지만, 회로당으로 환산하면 집합형 구조의 장점이 있다.",
            "반대로 우리 장치는 회로 단위 설치형이므로, 적은 회로부터 단계적으로 시작하기 쉽다는 장점이 있다.",
        ],
    ),
    (
        "5.6 Honeywell",
        [
            "Honeywell Forge Value+ 공식 페이지에서는 annual subscription이며 BMS point 수 기반 tier로 가격이 산정된다고 설명한다.",
            "즉 초기 구축비와 연간 운영비가 분리되는 전형적인 엔터프라이즈 SaaS 모델이다.",
            "우리에게 중요한 점은 상위 플랫폼 가격경쟁보다 현장 말단 데이터 장치의 필요성을 먼저 설득하는 것이다.",
        ],
    ),
]

for heading, bullets in vendor_details:
    add_heading(doc, heading, 2)
    add_bullets(doc, bullets)

add_heading(doc, "6. 우리 Smart Load 가격 포지셔닝 제안", 1)
add_bullets(doc, [
    "장치 단가만 비교하지 말고, 센서 수량 누적비용과 상위 플랫폼 운영비까지 포함해 경쟁사를 봐야 한다.",
    "우리 제품은 '분기회로 데이터 + 자체 차단 + DCU 연동'을 묶은 실행형 단말로 설명해야 한다.",
    "가격표는 1대 단가보다 '1회로 시작 가능', '단계적 확장 가능', '상위 시스템 교체 없이 붙일 수 있음'을 강조하는 방식이 유리하다.",
    "향후에는 장치 판매와 서버 분석/유지관리 서비스를 분리한 2단 가격모델도 검토할 수 있다.",
])

add_heading(doc, "7. 추천 표기 방식", 1)
add_bullets(doc, [
    "하드웨어 단가: Smart Load 본체, 통신부, 설치부자재 분리",
    "구축비: DCU 연동 세팅, 현장 셋업, 시험/보정 분리",
    "운영비: 서버 분석 또는 유지관리 기능이 붙을 경우 별도 표기",
    "고객 제안서에는 회로 수 증가에 따른 확장비용 구조를 함께 표시",
])

add_heading(doc, "8. 출처", 1)
sources = [
    "Schneider A9MEM1520 Newark: https://www.newark.com/schneider-electric/a9mem1520/energy-sensor-circuit-breaker/dp/84AH4718",
    "Schneider PowerTag Selection Guide: https://www.se.com/us/en/download/document/CA908058E/",
    "Siemens Sentron Powermind Product Sheet: https://assets.ctfassets.net/17si5cpawjzf/cwiIpmBdixlon9NxMKlsH/f801363861643d883b0e80f77511d3c5/App_SENTRONpowermind_ProductSheet_v2.0.pdf",
    "Eaton PXBCM-MB-BASIC PLATT: https://www.platt.com/p/1719348/eaton/c-h-pxbcm-mb-basic-pxbcm-mete/786687035824/cutpxbcmmbbasic",
    "ABB 2026 Thailand Price List snippet: https://library.e.abb.com/public/6ab0e443335144ed8f6705089a19f5e0/9AKK108472A0505_th_B_02122025%20All%20product%20pricelist%202026%20rev001.pdf",
    "ABB Emax 2 info: https://new.abb.com/low-voltage/products/circuit-breakers/emax2/stay-tuned1",
    "Leviton 71D48 official: https://leviton.com/products/71d48",
    "Leviton 71D48 public distributor price: https://lightingsupply.com/products/leviton-leviton-71d48",
    "Honeywell Forge Value+: https://buildings.honeywell.com/us/en/products/by-category/building-management/software/cloud-software/forge-value-plus",
    "LS ELECTRIC 회사 소개: https://www.ls-electric.com/ko/company/about/",
    "HD현대일렉트릭 INTEGRICT: https://www.hyundai-electric.com/elect/ko/integrict/integrict.jsp",
]
for s in sources:
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(9)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(OUT)
